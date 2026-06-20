from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PRD.md"
OUTPUT = ROOT / "docs" / "Pino-Production-System-PRD.docx"

NAVY = RGBColor(31, 55, 75)
TEAL = RGBColor(21, 122, 122)
GRAY = RGBColor(90, 98, 105)
LIGHT = "EAF3F3"


def set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pino Production System  |  ")
    set_font(run, 8.5, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, 9.5, color=NAVY)
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(part)
            set_font(run)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in (
    ("Heading 1", 16, NAVY, 14, 6),
    ("Heading 2", 12.5, TEAL, 10, 4),
    ("Heading 3", 11, NAVY, 8, 3),
):
    style = styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.text = "PRODUCT REQUIREMENTS DOCUMENT"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.runs[0], 8.5, bold=True, color=TEAL)
add_page_number(section.footer.paragraphs[0])

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(14)
title.paragraph_format.space_after = Pt(3)
run = title.add_run("PINO PRODUCTION SYSTEM")
set_font(run, 24, bold=True, color=NAVY)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
run = subtitle.add_run("Product Requirements Document")
set_font(run, 15, color=TEAL)

meta = doc.add_table(rows=5, cols=2)
meta.autofit = False
meta.columns[0].width = Inches(1.35)
meta.columns[1].width = Inches(5.45)
header_row = meta.rows[0]
header_row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
for cell, text in zip(header_row.cells, ("Document", "Details")):
    shade_cell(cell, "157A7A")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_font(p.add_run(text), 9.5, bold=True, color=RGBColor(255, 255, 255))
for row, (label, value) in zip(meta.rows[1:], (
    ("Status", "V1 baseline"),
    ("Version", "1.0 | 20 June 2026"),
    ("Product", "Responsive bilingual restaurant production and traceability platform"),
    ("Audience", "Product, engineering, QA, operations, and automated test-generation systems"),
)):
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
    shade_cell(row.cells[0], LIGHT)
    p = row.cells[0].paragraphs[0]
    p.clear()
    set_font(p.add_run(label), 9.5, bold=True, color=NAVY)
    p = row.cells[1].paragraphs[0]
    p.clear()
    set_font(p.add_run(value), 9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

lines = SOURCE.read_text(encoding="utf-8").splitlines()[7:]
for line in lines:
    stripped = line.strip()
    if not stripped:
        continue
    if stripped.startswith("# "):
        continue
    if stripped.startswith("## "):
        p = doc.add_paragraph(stripped[3:], style="Heading 1")
        continue
    if stripped.startswith("### "):
        p = doc.add_paragraph(stripped[4:], style="Heading 2")
        continue
    if re.match(r"^\d+\. ", stripped):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, re.sub(r"^\d+\. ", "", stripped))
        continue
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, stripped[2:])
        continue
    p = doc.add_paragraph()
    add_inline(p, stripped)

for p in doc.paragraphs:
    if p.style.name.startswith("List"):
        for run in p.runs:
            if not run.font.name:
                set_font(run)

props = doc.core_properties
props.title = "Pino Production System - Product Requirements Document"
props.subject = "Full V1 product requirements and test-generation baseline"
props.author = "Pino Production System"
props.keywords = "PRD, production, inventory, traceability, printing, testing"

doc.save(OUTPUT)
print(OUTPUT)
